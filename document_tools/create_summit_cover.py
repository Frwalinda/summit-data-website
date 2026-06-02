from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, JpegImagePlugin  # noqa: F401


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "summit_cover"
SOURCE_LOGO = Path(
    "/var/folders/xz/z3n4rnc56xxf20htjshd1d880000gn/T/TemporaryItems/"
    "NSIRD_screencaptureui_hbyoxj/Screenshot 2026-05-26 at 14.34.49.png"
)
CLEAN_LOGO = OUT_DIR / "summit_data_logo_clean.png"
ACCENT_RULE = OUT_DIR / "summit_accent_rule.png"
DOCX_OUT = OUT_DIR / "Summit_Commerce_OS_Cover_Page.docx"
PNG_OUT = OUT_DIR / "Summit_Commerce_OS_Cover_Page.png"
PDF_OUT = OUT_DIR / "Summit_Commerce_OS_Cover_Page.pdf"


NAVY = RGBColor(0x02, 0x1B, 0x2B)
TEAL = RGBColor(0x00, 0x9D, 0xC2)
STEEL = RGBColor(0x45, 0x58, 0x62)
MUTED = RGBColor(0x6A, 0x78, 0x80)


def set_run_font(run, name="Arial", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing


def paragraph_border_bottom(paragraph, color="009DC2", size="12", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_text(
    doc,
    text,
    *,
    size,
    color,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    before=0,
    after=0,
    line_spacing=None,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after, line_spacing=line_spacing)
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_spacer(doc, points):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=points)
    paragraph.add_run("")
    return paragraph


def clean_logo():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(SOURCE_LOGO).convert("RGBA")
    pixels = image.load()
    corners = [
        pixels[0, 0],
        pixels[image.width - 1, 0],
        pixels[0, image.height - 1],
        pixels[image.width - 1, image.height - 1],
    ]
    bg = tuple(sum(channel[i] for channel in corners) // len(corners) for i in range(3))

    threshold = 14
    for y in range(image.height):
        for x in range(image.width):
            r, g, b, a = pixels[x, y]
            distance = ((r - bg[0]) ** 2 + (g - bg[1]) ** 2 + (b - bg[2]) ** 2) ** 0.5
            if distance <= threshold and a > 0:
                pixels[x, y] = (255, 255, 255, 0)

    bbox = image.getbbox()
    if bbox:
        pad = 14
        left = max(0, bbox[0] - pad)
        top = max(0, bbox[1] - pad)
        right = min(image.width, bbox[2] + pad)
        bottom = min(image.height, bbox[3] + pad)
        image = image.crop((left, top, right, bottom))
    image.save(CLEAN_LOGO)


def make_accent_rule():
    width, height = 1200, 18
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    pixels = image.load()
    left = (0x00, 0x9D, 0xC2)
    right = (0x00, 0x42, 0x5F)
    for x in range(width):
        t = x / (width - 1)
        color = tuple(round(left[i] * (1 - t) + right[i] * t) for i in range(3))
        for y in range(height):
            alpha = 245 if 3 <= y <= 14 else 120
            pixels[x, y] = (*color, alpha)
    image.save(ACCENT_RULE)


def font(path, size):
    return ImageFont.truetype(str(path), size)


def draw_centered(draw, text, y, font_obj, fill, line_gap=12):
    lines = text.split("\n")
    current_y = y
    for line in lines:
        box = draw.textbbox((0, 0), line, font=font_obj)
        width = box[2] - box[0]
        height = box[3] - box[1]
        draw.text(((2550 - width) / 2, current_y), line, font=font_obj, fill=fill)
        current_y += height + line_gap
    return current_y


def build_preview_exports():
    page = Image.new("RGB", (2550, 3300), "white")
    draw = ImageDraw.Draw(page)

    arial = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
    arial_bold = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
    title_font = font(arial_bold, 120)
    subtitle_font = font(arial, 58)
    proposal_font = font(arial_bold, 69)
    footer_bold_font = font(arial_bold, 46)
    footer_font = font(arial, 46)

    logo = Image.open(CLEAN_LOGO).convert("RGBA")
    logo_w = 1155
    logo_h = round(logo.height * logo_w / logo.width)
    logo = logo.resize((logo_w, logo_h), Image.Resampling.LANCZOS)
    page.paste(logo, ((2550 - logo_w) // 2, 310), logo)

    y = 310 + logo_h + 155
    y = draw_centered(draw, "SUMMIT COMMERCE OS", y, title_font, (2, 27, 43), line_gap=0)

    rule_w = 825
    rule_y = y + 45
    for x in range(rule_w):
        t = x / max(1, rule_w - 1)
        left = (0, 157, 194)
        right = (0, 66, 95)
        color = tuple(round(left[i] * (1 - t) + right[i] * t) for i in range(3))
        draw.line(
            ((2550 - rule_w) // 2 + x, rule_y, (2550 - rule_w) // 2 + x, rule_y + 13),
            fill=color,
        )

    y = rule_y + 100
    y = draw_centered(
        draw,
        "Digital Commerce & Field Operations Platform\nfor Distributors, FMCGs and SMEs",
        y,
        subtitle_font,
        (69, 88, 98),
        line_gap=24,
    )

    y += 110
    y = draw_centered(
        draw,
        "Investor & Banking Partnership Proposal",
        y,
        proposal_font,
        (2, 27, 43),
        line_gap=0,
    )
    draw.line((785, y + 36, 1765, y + 36), fill=(212, 229, 234), width=4)

    draw_centered(draw, "Confidential", 2880, footer_bold_font, (106, 120, 128), line_gap=0)
    draw_centered(draw, "May 2026", 2945, footer_font, (106, 120, 128), line_gap=0)

    page.save(PNG_OUT, quality=95)
    page.save(PDF_OUT, resolution=300.0)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333


def build_docx():
    clean_logo()
    make_accent_rule()
    build_preview_exports()

    doc = Document()
    configure_document(doc)

    add_spacer(doc, 24)

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(logo_paragraph, after=20)
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(str(CLEAN_LOGO), width=Inches(3.85))

    add_text(
        doc,
        "SUMMIT COMMERCE OS",
        size=29,
        color=NAVY,
        bold=True,
        after=9,
        line_spacing=1.0,
    )

    rule_paragraph = doc.add_paragraph()
    rule_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(rule_paragraph, after=22)
    rule_run = rule_paragraph.add_run()
    rule_run.add_picture(str(ACCENT_RULE), width=Inches(2.75))

    add_text(
        doc,
        "Digital Commerce & Field Operations Platform\nfor Distributors, FMCGs and SMEs",
        size=14,
        color=STEEL,
        after=30,
        line_spacing=1.18,
    )

    proposal = add_text(
        doc,
        "Investor & Banking Partnership Proposal",
        size=16.5,
        color=NAVY,
        bold=True,
        after=5,
        line_spacing=1.05,
    )
    paragraph_border_bottom(proposal, color="D4E5EA", size="4", space="7")

    add_spacer(doc, 168)

    add_text(
        doc,
        "Confidential",
        size=11,
        color=MUTED,
        bold=True,
        after=2,
        line_spacing=1.0,
    )
    add_text(
        doc,
        "May 2026",
        size=11,
        color=MUTED,
        bold=False,
        after=0,
        line_spacing=1.0,
    )

    doc.save(DOCX_OUT)
    return DOCX_OUT


if __name__ == "__main__":
    print(build_docx())
    print(PNG_OUT)
    print(PDF_OUT)
